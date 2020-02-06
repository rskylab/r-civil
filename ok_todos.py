#!/usr/bin/env python
# coding: utf-8

# In[ ]:
import os
import docx
import openpyxl
import datetime
from docx.shared import Pt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from termcolor import cprint

diretorio = 'Z:/Mexico/NUDIVERSIS NUCORA/=MODELOS e docs para ATENDIMENTO=/.REQUALIFICAÇÃO CIVIL/MODELOS DE OFÍCIOS PARA CERTIDÕES NEGATIVAS/'

ctrl_p = 'Z:/Mexico/NUDIVERSIS 2020/CONTROLES/Ofícios NUDIVERSIS 2020/Ofícios Requalificação 2020.xlsx'
controle = openpyxl.load_workbook(filename = ctrl_p) 
cs = controle['Plan1']

dia = datetime.date.today().day
mes = datetime.date.today().month
ano = datetime.date.today().year

assistido = input('Qual o nome dx assistidx? ')
sl_assistido = assistido[0:31]

sn = input('O nome do assistido está correto? ')
while 'o' in sn:
    assistido = input('Qual o nome dx assistidx? ')
    sn = input('O nome do assistido está correto? ')

print('\n')

lista_est = ['1', 'Mariana', 'mariana', '2', 'Melissa', 'melissa', '3', 'Wagner', 'wagner', '4', 'Thiago', 'thiago']

pergunta_estagiarios = 'Que estagiário está fazendo a requalificação? \n1. Mariana \n2. Melissa \n3. Wagner \n4. Thiago \n\n'

quem = input(pergunta_estagiarios)

while quem not in lista_est:
    print(f'{quem} não é reconhecido como um dos estagiários. Digite novamente.')
    quem = input(pergunta_estagiarios)
    
if quem == '1' or quem == 'Mariana' or quem == 'mariana':
    plan_p = 'Atendimento.xlsx'
    os.chdir('Z:/Mexico/NUDIVERSIS 2020/Estagiários/Mariana/Requalificação civil')
    estagiario = 'Mariana Stillner'
    
if quem == '2' or quem == 'Melissa' or quem == 'melissa':
    plan_p = 'Atendimento.xlsx'
    os.chdir('Z:/Mexico/NUDIVERSIS 2020/Estagiários/Melissa/NUDIVERSIS/REQUALIFICAÇÃO CIVIL')
    estagiario = 'Melissa Kreil'
    
if quem == '3' or quem == 'Wagner' or quem == 'wagner':
    plan_p = 'Atendimento.xlsx'
    os.chdir('Z:/Mexico/NUDIVERSIS 2020/Estagiários/Wagner Rabelo/NUDIVERSIS')
    estagiario = 'Wagner Rabelo'
    
if quem == '4' or quem == 'Thiago' or quem == 'thiago':
    plan_p = 'Atendimento.xlsx'
    os.chdir('Z:/Mexico/NUDIVERSIS 2020/Estagiários/Thiago Percides/NUDIVERSIS/Requalificação Civil')
    estagiario = 'Thiago Percides'


wb = openpyxl.load_workbook(filename = plan_p)

def criar_planilha():
    
    p_nome = os.path.basename(plan_p)
    arq = os.path.splitext(p_nome)[0]
        
    if sl_assistido in wb.sheetnames: 
        if not os.path.exists(assistido):
            os.mkdir(assistido)
        print('Já existe uma planilha com o nome do assistido. Preencha com os dados e siga para a próxima célula.\n\nCaso a pasta não tenha sido criada, ou seja outro atendimento, clique no botão "Kernel" acima, e depois em "Restart". Depois de reiniciar, rode esta célula novamente.')

    else:
        try:
            wb.save(plan_p)
        except PermissionError:
            print(f'Feche a planilha {arq} e rode esta célula novamente.')
            return
                  
        ws = wb.create_sheet(sl_assistido)
        ws['A1'] = 'Nome'
        ws['B1'] = 'Nome de registro'
        ws['C1'] = 'Nacionalidade'
        ws['D1'] = 'Estado civil'
        ws['E1'] = 'Ocupação'
        ws['F1'] = 'N° da identidade' 
        ws['G1'] = 'Órgão expedidor'
        ws['H1'] = 'CPF' 
        ws['I1'] = 'Endereço completo'
        ws['J1'] = 'Bairro'
        ws['K1'] = 'Cidade' 
        ws['L1'] = 'CEP'
        ws['M1'] = 'Telefone'
        ws['N1'] = 'Cidades p/ gerar ofício'

        ws['A2'] = assistido

        ws.column_dimensions['A'].width = 25
        ws.column_dimensions['B'].width = 25
        ws.column_dimensions['C'].width = 13
        ws.column_dimensions['D'].width = 13
        ws.column_dimensions['E'].width = 15
        ws.column_dimensions['F'].width = 17
        ws.column_dimensions['G'].width = 25
        ws.column_dimensions['H'].width = 17
        ws.column_dimensions['I'].width = 30
        ws.column_dimensions['J'].width = 15
        ws.column_dimensions['K'].width = 15
        ws.column_dimensions['L'].width = 17
        ws.column_dimensions['M'].width = 15
        ws.column_dimensions['N'].width = 25

        wb.save(plan_p)
        
        if not os.path.exists(assistido):
            os.mkdir(assistido)
        
        print(f'Foi criada a planilha {assistido} no arquivo {arq}.')
        return

lista_oficios = []

jmp = diretorio + 'Padrão/Justiça Militar.docx'
jm = Document(jmp)

trtp = diretorio + 'Padrão/TRT.docx'
trt = Document(trtp)

padrao_doc = [jm, trt]
padrao_path = [jmp, trtp]
padrao_zip = zip(padrao_doc, padrao_path)

riod1p = diretorio + 'Rio de Janeiro/Distribuidor (Capital) - 1°.docx'
riod1 = Document(riod1p)

riod2p = diretorio + 'Rio de Janeiro/Distribuidor (Capital) - 2°.docx'
riod2 = Document(riod2p)

riod3p = diretorio + 'Rio de Janeiro/Distribuidor (Capital) - 3°.docx'
riod3 = Document(riod3p)

riod4p = diretorio + 'Rio de Janeiro/Distribuidor (Capital) - 4°.docx'
riod4 = Document(riod4p)

riod5p = diretorio + 'Rio de Janeiro/Distribuidor (Capital) - 5°.docx'
riod5 = Document(riod5p)

riod6p = diretorio + 'Rio de Janeiro/Distribuidor (Capital) - 6°.docx'
riod6 = Document(riod6p)

riod7p = diretorio + 'Rio de Janeiro/Distribuidor (Capital) - 7°.docx'
riod7 = Document(riod7p)

riod8p = diretorio + 'Rio de Janeiro/Distribuidor (Capital) - 8°.docx'
riod8= Document(riod8p)

riod9p = diretorio + 'Rio de Janeiro/Distribuidor (Capital) - 9°.docx'
riod9 = Document(riod9p)

rioi1p = diretorio + 'Rio de Janeiro/Interdição e Tutelas - 1°.docx' 
rioi1 = Document(rioi1p)

rioi2p = diretorio + 'Rio de Janeiro/Interdição e Tutelas - 2°.docx' 
rioi2 = Document(rioi2p)

riot1p = diretorio + 'Rio de Janeiro/Títulos e Protestos - 1°.docx' 
riot1 = Document(riot1p)

riot2p = diretorio + 'Rio de Janeiro/Títulos e Protestos - 2°.docx'
riot2= Document(riot2p)

riot3p = diretorio + 'Rio de Janeiro/Títulos e Protestos - 3°.docx'
riot3 = Document(riot3p)

riot4p = diretorio + 'Rio de Janeiro/Títulos e Protestos - 4°.docx'
riot4 = Document(riot4p)

of_rio = [riod1, riod2, riod3, riod4, riod5, riod6, riod7, riod8, riod9, rioi1, rioi2, riot1, riot2, riot3, riot4]
p_rio = [riod1p, riod2p, riod3p, riod4p, riod5p, riod6p, riod7p, riod8p, riod9p, rioi1p, rioi2p, riot1p, riot2p, riot3p, riot4p]
z_rio = zip(of_rio, p_rio)

######

brdp = diretorio + 'Belford Roxo/Distribuidor (Belford Roxo).docx'
brd = Document(brdp)

brtp = diretorio + 'Belford Roxo/Títulos e Protestos (Belford Roxo).docx'
brt = Document(brtp)

brip = diretorio + 'Belford Roxo/Interdição e Tutelas (Belford Roxo).docx'
bri = Document(brip)

of_br = [brd, brt, bri]
p_br = [brdp, brtp, brip]
z_br = zip(of_br, p_br)

######

dcd1p = diretorio + 'Duque de Caxias/Distribuidor (Duque de Caxias) - 1°.docx' 
dcd1 = Document(dcd1p)

dcd2p = diretorio + 'Duque de Caxias/Distribuidor (Duque de Caxias) - 2°.docx' 
dcd2 = Document(dcd2p)

dcip = diretorio + 'Duque de Caxias/Interdição e Tutelas (Caxias).docx' 
dci = Document(dcip)

dct1p = diretorio + 'Duque de Caxias/Títulos e Protestos (Caxias) - 1°.docx' 
dct1 = Document(dct1p)

dct2p = diretorio + 'Duque de Caxias/Títulos e Protestos (Caxias) - 2°.docx' 
dct2 = Document(dct2p)

of_dc = [dcd1, dcd2, dci, dct1, dct2]
p_dc = [dcd1p, dcd2p, dcip, dct1p, dct2p]
z_dc = zip(of_dc, p_dc)

######

madp = diretorio + 'Magé/Distribuidor (Magé).docx' 
mad = Document(madp)

matp = diretorio + 'Magé/Títulos e Protestos (Magé).docx' 
mat = Document(matp)

maip = diretorio + 'Magé/Interdição e Tutelas (Magé).docx' 
mai = Document(maip)

of_ma = [mad, mat, mai]
p_ma = [madp, matp, maip]
z_ma = zip(of_ma, p_ma)

######

ntd1p = diretorio + 'Niterói/Distribuidor (Niterói) - 1°.docx' 
ntd1 = Document(ntd1p)

ntd3p = diretorio + 'Niterói/Distribuidor (Niterói) - 3°.docx' 
ntd3 = Document(ntd3p)

ntip = diretorio + 'Niterói/Interdição e Tutelas (Niterói).docx' 
nti = Document(ntip)

ntt1p = diretorio + 'Niterói/Títulos e Protestos (Niterói) - 1°.docx' 
ntt1 = Document(ntt1p)

ntt2p = diretorio + 'Niterói/Títulos e Protestos (Niterói) - 2°.docx' 
ntt2 = Document(ntt2p)

ntt13p = diretorio + 'Niterói/Títulos e Protestos (Niterói) - 13°.docx' 
ntt13 = Document(ntt13p)

of_nt = [ntd1, ntd3, nti, ntt1, ntt2, ntt13]
p_nt = [ntd1p, ntd3p, ntip, ntt1p, ntt2p, ntt13p]
z_nt = zip(of_nt, p_nt)

######

nidp = diretorio + 'Nova Iguaçu/Distribuidor (Nova Iguaçu).docx' 
nid = Document(nidp)

nitp = diretorio + 'Nova Iguaçu/Títulos e Protestos (Nova Iguaçu).docx' 
nit = Document(nitp)

niip = diretorio + 'Nova Iguaçu/Interdição e Tutelas (Nova Iguaçu).docx' 
nii = Document(niip)

of_ni = [nid, nit, nii]
p_ni = [nidp, nitp, niip]
z_ni = zip(of_ni, p_ni)

######

sgdp = diretorio + 'São Gonçalo/Distribuidor (São Gonçalo).docx' 
sgd = Document(sgdp)

sgtp = diretorio + 'São Gonçalo/Títulos e Protestos (São Gonçalo).docx' 
sgt= Document(sgtp)

sgip = diretorio + 'São Gonçalo/Interdição e Tutelas (São Gonçalo).docx'
sgi = Document(sgip)

of_sg = [sgd, sgt, sgi]
p_sg = [sgdp, sgtp, sgip]
z_sg = zip(of_sg, p_sg)

######

sjmdp = diretorio + 'São João de Meriti/Distribuidor (São João de Meriti).docx' 
sjmd= Document(sjmdp)

sjmtp = diretorio + 'São João de Meriti/Títulos e Protestos (São João de Meriti).docx'
sjmt = Document(sjmtp)

sjmip = diretorio + 'São João de Meriti/Interdição e Tutelas (São João de Meriti).docx' 
sjmi = Document(sjmip)

of_sjm = [sjmd, sjmt, sjmi]
p_sjm = [sjmdp, sjmtp, sjmip]
z_sjm = zip(of_sjm, p_sjm)

######

srdp = diretorio + 'Seropédica/Distribuidor (Seropédica).docx' 
srd = Document(srdp)

srtp = diretorio + 'Seropédica/Títulos e Protestos (Seropédica).docx' 
srt = Document(srtp)

srip = diretorio + 'Seropédica/Interdição e Tutelas (Seropédica).docx' 
sri = Document(srip)

of_sr = [srd, srt, sri]
p_sr = [srdp, srtp, srip]
z_sr = zip(of_sr, p_sr)

#####

def gerar_oficios():

    wb = openpyxl.load_workbook(filename = plan_p) 
    ws = wb[sl_assistido]
    
    if not os.path.exists(assistido):
        os.mkdir(assistido)
    #assistido = ws['A2'].value
    
    telefone = str(ws['M2'].value)
    
    none_count = 0
    for i in ws.iter_rows(min_row=2, max_row=2, min_col = 1, max_col=14, values_only=True):
        for x in i:
            if x == None:
                none_count +=1
            
    if none_count != 0:
        print(f'Há {none_count} colunas vazias na planilha {assistido}. Preencha e salve a planilha novamente.')
        return
    
    dados = f', cujo nome de registro é <<<nome de registro>>>, <<<nacionalidade>>>, <<<estado civil>>>, <<<ocupação>>>, IDENTIDADE n.º <<<identidade>>>, <<<órgão expedidor>>>, CPF: n.º <<<cpf>>>, residente na <<<endereço>>>, <<<bairro>>>, <<<cidade>>>, CEP: <<<cep>>>, '
    dados = dados.replace('<<<nome de registro>>>', ws['B2'].value)
    dados = dados.replace('<<<nacionalidade>>>', ws['C2'].value)
    dados = dados.replace('<<<estado civil>>>', ws['D2'].value)
    dados = dados.replace('<<<ocupação>>>', ws['E2'].value)
    dados = dados.replace('<<<identidade>>>', str(ws['F2'].value))
    dados = dados.replace('<<<órgão expedidor>>>', ws['G2'].value)
    dados = dados.replace('<<<cpf>>>', str(ws['H2'].value))
    dados = dados.replace('<<<endereço>>>', ws['I2'].value)
    dados = dados.replace('<<<bairro>>>', ws['J2'].value)
    dados = dados.replace('<<<cidade>>>', ws['K2'].value)
    dados = dados.replace('<<<cep>>>', str(ws['L2'].value))
    
    lista_cidades = []
    for i in ws.iter_rows(min_row=2, max_row=30, min_col = 14, max_col = 14, values_only=True):
        if i[0] != None:
            lista_cidades.append(i[0])

    lista_erros = ['Rio de Janeiro ','rio de janeiro', 'rio de janeiro ', ' Rio de Janeiro', ' rio de janeiro', 'Rio de janeiro', ' Rio de janeiro', 'Rio de janeiro ', 'Belford Roxo ', ' Belford Roxo', 'belford roxo', ' belford roxo', 'belford roxo ', 'Belford roxo', ' Belford roxo', 'Belford roxo ', 'Duque de Caxias ', ' Duque de Caxias', 'duque de caxias', ' duque de caxias', 'duque de caxias ', 'Duque de caxias', ' Duque de caxias', 'Duque de caxias ', ' Magé', 'Magé ', 'magé', 'magé ', ' magé', 'Mage', 'Mage ', ' Mage', 'mage', ' mage', 'mage ', 'Niterói ', ' Niterói', 'niterói', ' niterói', 'niterói ', 'Niteroi', 'Niteroi ', ' Niteroi', 'niteroi', 'niteroi ', ' niteroi', 'Nova Iguaçu ', ' Nova Iguaçu', 'Nova iguaçu', 'Nova iguaçu ', ' Nova iguaçu', 'nova iguaçu', 'nova iguaçu ', ' nova iguaçu', 'São Gonçalo ', ' São Gonçalo', 'São gonçalo', ' São gonçalo', 'São gonçalo ', 'são gonçalo', 'são gonçalo ', ' são gonçalo', 'Sao Gonçalo', ' Sao Gonçalo', 'Sao Gonçalo ', 'São João de Meriti ', ' São João de Meriti', ' São joão de meriti', 'São joão de meriti', 'São joão de meriti ', 'são joão de meriti', ' são joão de meriti', 'são joão de meriti ', ' Seropédica', 'Seropédica ', 'seropédica', ' seropédica', 'seropédica ', 'Seropedica', ' Seropedica', 'Seropedica ', 'seropedica', ' seropedica', 'seropedica ']
                   
    subfolders = [ f.name for f in os.scandir(diretorio) if f.is_dir() ]
    for i in lista_cidades:
        if i not in subfolders:
            if i not in lista_erros:
                print(f'A célula {i} não foi reconhecida como uma cidade da lista de modelos. A pasta de modelos tem as seguintes sub-pastas: {subfolders}.\n\nCaso a cidade ESTEJA na lista, digite novamente, e salve a planilha de novo. Cuidado com erros de digitação.\n\nDo contrário, salve a planilha apenas com as cidades para as quais há modelos, e rode a célula abaixo para gerar os demais ofícios manualmente.')
                return
    
    lista_oficios.append(padrao_zip)

    if 'Rio de Janeiro' in lista_cidades or 'Rio de Janeiro ' in lista_cidades or 'rio de janeiro' in lista_cidades or 'rio de janeiro ' in lista_cidades or ' Rio de Janeiro' in lista_cidades or ' rio de janeiro' in lista_cidades or 'Rio de janeiro' in lista_cidades or ' Rio de janeiro' in lista_cidades or 'Rio de janeiro ' in lista_cidades:

        lista_oficios.append(z_rio)

    if 'Belford Roxo' in lista_cidades or 'Belford Roxo ' in lista_cidades or ' Belford Roxo' in lista_cidades or 'belford roxo' in lista_cidades or ' belford roxo' in lista_cidades or 'belford roxo ' in lista_cidades or 'Belford roxo' in lista_cidades or ' Belford roxo' in lista_cidades or 'Belford roxo ' in lista_cidades:

        lista_oficios.append(z_br)

    if 'Duque de Caxias' in lista_cidades or 'Duque de Caxias ' in lista_cidades or ' Duque de Caxias' in lista_cidades or 'duque de caxias' in lista_cidades or ' duque de caxias' in lista_cidades or 'duque de caxias ' in lista_cidades or 'Duque de caxias' in lista_cidades or ' Duque de caxias' in lista_cidades or 'Duque de caxias ' in lista_cidades:

        lista_oficios.append(z_dc)

    if 'Magé' in lista_cidades or ' Magé' in lista_cidades or 'Magé ' in lista_cidades or 'magé' in lista_cidades or 'magé ' in lista_cidades or ' magé' in lista_cidades or 'Mage' in lista_cidades or 'Mage ' in lista_cidades or ' Mage' in lista_cidades or 'mage' in lista_cidades or ' mage' in lista_cidades or 'mage ' in lista_cidades:

        lista_oficios.append(z_ma)

    if 'Niterói' in lista_cidades or 'Niterói ' in lista_cidades or ' Niterói' in lista_cidades or 'niterói' in lista_cidades or ' niterói' in lista_cidades or 'niterói ' in lista_cidades or 'Niteroi' in lista_cidades or 'Niteroi ' in lista_cidades or ' Niteroi' in lista_cidades or 'niteroi' in lista_cidades or 'niteroi ' in lista_cidades or ' niteroi' in lista_cidades:

        lista_oficios.append(z_nt)

    if 'Nova Iguaçu' in lista_cidades or 'Nova Iguaçu ' in lista_cidades or ' Nova Iguaçu' in lista_cidades or 'Nova iguaçu' in lista_cidades or 'Nova iguaçu ' in lista_cidades or ' Nova iguaçu' in lista_cidades or 'nova iguaçu' in lista_cidades or 'nova iguaçu ' in lista_cidades or ' nova iguaçu' in lista_cidades:

        lista_oficios.append(z_ni)

    if 'São Gonçalo' in lista_cidades or 'São Gonçalo ' in lista_cidades or ' São Gonçalo' in lista_cidades or 'São gonçalo' in lista_cidades or ' São gonçalo' in lista_cidades or 'São gonçalo ' in lista_cidades or 'são gonçalo' in lista_cidades or 'são gonçalo ' in lista_cidades or ' são gonçalo' in lista_cidades or 'Sao Gonçalo' in lista_cidades or ' Sao Gonçalo' in lista_cidades or 'Sao Gonçalo ' in lista_cidades:

        lista_oficios.append(z_sg)

    if 'São João de Meriti' in lista_cidades or 'São João de Meriti ' in lista_cidades or ' São João de Meriti' in lista_cidades or ' São joão de meriti' in lista_cidades or 'São joão de meriti' in lista_cidades or 'São joão de meriti ' in lista_cidades or 'são joão de meriti' in lista_cidades or ' são joão de meriti' in lista_cidades or 'são joão de meriti ' in lista_cidades:

        lista_oficios.append(z_sjm)

    if 'Seropédica' in lista_cidades or ' Seropédica' in lista_cidades or 'Seropédica ' in lista_cidades or 'seropédica' in lista_cidades or ' seropédica' in lista_cidades or 'seropédica ' in lista_cidades or 'Seropedica' in lista_cidades or ' Seropedica' in lista_cidades or 'Seropedica ' in lista_cidades or 'seropedica' in lista_cidades or ' seropedica' in lista_cidades or 'seropedica ' in lista_cidades:

        lista_oficios.append(z_sr)
    
    qtd_of = 0
    
    req = Document(diretorio + 'Padrão/Requerimento.docx')
    
    req.paragraphs[11].text = ''
    aa = req.paragraphs[11]
    um = aa.add_run(assistido)
    um.bold = True 
    um.font.name = 'Arial'
    um.font.size = Pt(11)

    dois = aa.add_run(dados)
    dois.font.name = 'Arial'
    dois.font.size = Pt(11)
    
    tres = aa.add_run('telefone: ' + telefone + ' ')
    tres.font.name = 'Arial'
    tres.font.size = Pt(11)

    req.save(f'./{assistido}/Requerimento TRT.docx')
    
    anx = Document(diretorio + 'Padrão/Anexo 1° e 2° Grau.docx')
    anx.save(f'./{assistido}/Anexo 1° e 2° Grau.docx')
   
    for i in lista_oficios:
        for x in i:
            try:
                controle.save(ctrl_p)
            except PermissionError:
                print('A planilha do controle de ofícios está aberta. Feche o controle e rode esta célula novamente.')
                return
          
            doc = x[0]
            caminho = x[1]

            n_atual = cs['F2'].value
            numero = n_atual + 1
            tabela = doc.tables[0]
            t1 = tabela.rows[0].cells[0]
            t2 = tabela.rows[0].cells[1]
            t1.text = ''
            p = t1.add_paragraph(f'Ofício n° {numero}/REQ/{ano}/NUDIVERSIS/DPGE/RJ')
            t2.text = ''
            q = t2.add_paragraph(f'Rio de Janeiro, {dia}/{mes}/{ano}')
            q.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            for row in tabela.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(11)
            
            doc.paragraphs[11].text = ''
            aa = doc.paragraphs[11]
            um = aa.add_run(assistido)
            um.bold = True 
            um.font.name = 'Arial'
            um.font.size = Pt(11)

            dois = aa.add_run(dados)
            dois.font.name = 'Arial'
            dois.font.size = Pt(11)

            doisemeio = aa.add_run('com a finalidade de ')
            doisemeio.font.name = 'Arial'
            doisemeio.font.size = Pt(11)
            
            tres = aa.add_run('INSTRUIR PROCESSO ADMINISTRATIVO para retificação dos assentamentos civis.')
            tres.font.name = 'Arial'
            tres.font.size = Pt(11)
            tres.bold = True
            tres.underline = True

            base = os.path.basename(caminho)
            cartorio = os.path.splitext(base)[0]

            cs.insert_rows(2)
            cs['F2'] = numero
            cs['G2'] = cartorio
            cs['H2'] = assistido
            cs['I2'] = 'Certidão Negativa'
            cs['J2'] = estagiario

            doc.save(f'./{assistido}/Ofício Nudiversis n. {numero}.Req.{ano} - ' + cs['G2'].value + f' - {assistido}.docx')
            controle.save(ctrl_p)
            qtd_of += 1
            
    print(f'Foram gerados {qtd_of} ofícios na pasta {assistido}. \n')
    cprint('Lembre-se de mudar o PRENOME e o SEXO no documento Requerimento TRT.', attrs = ['bold', 'underline'])
    return


def oficios_extra():
    
    if not os.path.exists(assistido):
        os.mkdir(assistido)
            
    wb = openpyxl.load_workbook(filename = plan_p)
    ws = wb[sl_assistido]
      
    quantos = int(input('De quantos ofícios você precisa? '))
    qtd_of = 0
    
    doc = Document(diretorio + 'Padrão/Modelo padrão.docx')
    
    dados = f', cujo nome de registro é <<<nome de registro>>>, <<<nacionalidade>>>, <<<estado civil>>>, <<<ocupação>>>, IDENTIDADE n.º <<<identidade>>>, <<<órgão expedidor>>>, CPF: n.º <<<cpf>>>, residente na <<<endereço>>>, <<<bairro>>>, <<<cidade>>>, CEP: <<<cep>>>, '
    dados = dados.replace('<<<nome de registro>>>', ws['B2'].value)
    dados = dados.replace('<<<nacionalidade>>>', ws['C2'].value)
    dados = dados.replace('<<<estado civil>>>', ws['D2'].value)
    dados = dados.replace('<<<ocupação>>>', ws['E2'].value)
    dados = dados.replace('<<<identidade>>>', str(ws['F2'].value))
    dados = dados.replace('<<<órgão expedidor>>>', ws['G2'].value)
    dados = dados.replace('<<<cpf>>>', str(ws['H2'].value))
    dados = dados.replace('<<<endereço>>>', ws['I2'].value)
    dados = dados.replace('<<<bairro>>>', ws['J2'].value)
    dados = dados.replace('<<<cidade>>>', ws['K2'].value)
    dados = dados.replace('<<<cep>>>', str(ws['L2'].value))
    
    while quantos != 0:
        n_atual = cs['F2'].value
        numero = n_atual + 1
        tabela = doc.tables[0]
        t1 = tabela.rows[0].cells[0]
        t2 = tabela.rows[0].cells[1]
        t1.text = ''
        p = t1.add_paragraph(f'Ofício n° {numero}/REQ/{ano}/NUDIVERSIS/DPGE/RJ')
        t2.text = ''
        q = t2.add_paragraph(f'Rio de Janeiro, {dia}/{mes}/{ano}')
        q.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        for row in tabela.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)

        doc.paragraphs[11].text = ''
        aa = doc.paragraphs[11]
        um = aa.add_run(assistido)
        um.bold = True 
        um.font.name = 'Arial'
        um.font.size = Pt(11)

        dois = aa.add_run(dados)
        dois.font.name = 'Arial'
        dois.font.size = Pt(11)

        doisemeio = aa.add_run('com a finalidade de ')
        doisemeio.font.name = 'Arial'
        doisemeio.font.size = Pt(11)

        tres = aa.add_run('INSTRUIR PROCESSO ADMINISTRATIVO para retificação dos assentamentos civis.')
        tres.font.name = 'Arial'
        tres.font.size = Pt(11)
        tres.bold = True
        tres.underline = True

        cs.insert_rows(2)
        cs['F2'] = numero
        cs['H2'] = assistido
        cs['I2'] = 'Certidão Negativa'
        cs['J2'] = estagiario

        doc.save(f'./{assistido}/Ofício Nudiversis n. {numero}.Req.{ano} - {assistido}.docx')
        controle.save(ctrl_p)
        qtd_of += 1
        quantos -= 1
            
    print(f'\nForam gerados {qtd_of} ofícios adicionais na pasta {assistido}.')
    cprint('Lembre-se de preencher os documentos com os cartórios e atualizá-los na planilha.', attrs = ['bold', 'underline'])
    return
    